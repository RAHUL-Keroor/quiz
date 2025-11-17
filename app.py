import os
import io
import re
import ast
import uuid
import random
from datetime import datetime, timedelta

import pytz
import PyPDF2
from bson import ObjectId
from pymongo import MongoClient
from flask import (
    Flask, render_template, request, redirect,
    url_for, session, jsonify, send_file
)
from flask_mail import Mail, Message
from fpdf import FPDF
from docx import Document
from werkzeug.utils import secure_filename

import tempfile


# --------------------------------------------------------------------------
# --- 1. Configuration & App Initialization ---

app = Flask(__name__)
app.secret_key = 'super-secret-key-change-this-in-production'
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'rahulkeroor6@gmail.com'      # change this
app.config['MAIL_PASSWORD'] = 'arqngatvmalvbxhc'
app.config['MAIL_DEFAULT_SENDER'] = app.config['MAIL_USERNAME']       # NOT your Gmail password
mail = Mail(app)




@app.context_processor
def inject_datetime():
    return dict(datetime=datetime)

app.secret_key = 'super-secret-key-change-this-in-production'
DATABASE = 'quiz.db'
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)



# --------------------------------------------------------------------------
# --- 2. Database Utility Functions ---
# --------------------------------------------------------------------------


# Load SpaCy model for NLP question gene
from pymongo import MongoClient

def get_db_connection():
    mongo_uri = os.getenv("MONGO_URI")
    client = MongoClient(mongo_uri)
    return client["quiz_app_db"]


# Helper functions
def options_to_db(options_list):
    return ";;;".join(options_list)

def options_from_db(options_str):
    return options_str.split(";;;")

# --------------------------------------------------------------------------
# --- 3. Improved NLP-Powered Question Generator ---


import re
import random
from PyPDF2 import PdfReader

def generate_questions_from_pdf(pdf_path, num_questions=5):
    """
    Simple + Render-friendly PDF question generator.
    Creates:
      - MCQ questions
      - Fill in the blanks questions
    """
    text = ""

    # Extract text safely
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + " "
    except:
        print("❌ PDF read failed — returning empty list")
        return []

    # Clean text
    text = re.sub(r'\s+', ' ', text)

    # Split into simple sentences
    sentences = [s.strip() for s in re.split(r'[.!?]', text) if len(s.split()) > 4]

    if not sentences:
        return []

    questions = []

    for i in range(min(num_questions, len(sentences))):
        sent = sentences[i]

        # Choose a random word for blank
        words = [w for w in sent.split() if len(w) > 4]
        if not words:
            continue

        correct = random.choice(words)

        # ------------- Fill in the blanks -------------
        qtext_fill = sent.replace(correct, "_____")

        # ------------- MCQ distractors -------------
        distractors = random.sample(words, min(3, len(words)))
        if correct in distractors:
            distractors.remove(correct)

        while len(distractors) < 3:
            distractors.append("Option" + str(random.randint(1, 100)))

        options = [correct] + distractors[:3]
        random.shuffle(options)

        questions.append({
            "question_text": qtext_fill,
            "options": options,
            "correct_answer": correct,
            "description": ""
        })

    return questions


@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    pdf = request.files.get('pdf_file')
    num_questions = int(request.form.get('num_questions', 5))
    Quiz_name = request.form.get('Quiz_name', 'Anonymous')
    timer = int(request.form.get('timer_minutes', 5))
    is_shuffled = request.form.get('is_shuffled') == 'true'

    usn_start = request.form.get('usn_start', '').upper()
    usn_end = request.form.get('usn_end', '').upper()

    questions = []

    # CASE 1 — No PDF
    if not pdf or pdf.filename.strip() == "":
        print("⚠ No PDF uploaded")
        for i in range(num_questions):
            questions.append({
                "question_text": f"Question {i+1}: ",
                "options": ["", "", "", ""],
                "correct_answer": "",
                "description": ""
            })

    # CASE 2 — PDF provided
    else:
        TEMP_DIR = tempfile.gettempdir()
        filename = secure_filename(pdf.filename)
        save_path = os.path.join(TEMP_DIR, filename)

        pdf.save(save_path)

        try:
            questions = generate_questions_from_pdf(save_path, num_questions)
        except:
            print("❌ PDF error, fallback to empty questions")
            for i in range(num_questions):
                questions.append({
                    "question_text": f"Question {i+1}: ",
                    "options": ["", "", "", ""],
                    "correct_answer": "",
                    "description": ""
                })

    return render_template(
        'display_generated_questions.html',
        questions=questions,
        num_questions=len(questions),
        Quiz_name=Quiz_name,
        timer=timer,
        is_shuffled=is_shuffled,
        usn_start=usn_start,
        usn_end=usn_end,
        start_time=request.form.get('start_time')
    )



@app.route('/save_quiz', methods=['POST'])
def save_quiz():
    num_questions = int(request.form.get('num_questions', 0))
    Quiz_name = request.form.get('Quiz_name', 'Anonymous')

    if num_questions == 0:
        return "No questions received", 400

    quiz_code = str(uuid.uuid4())[:8].upper()
    creator_token = str(random.randint(100000, 999999))
    quiz_link = f"/results/{quiz_code}?token={creator_token}"

    db = get_db_connection()
    india_tz = pytz.timezone("Asia/Kolkata")
    start_time_str = request.form.get('start_time')

    print(f"🕒 Received from form: {start_time_str}")

    if start_time_str:
        try:
            today = datetime.now(india_tz).date()
            hours, minutes = map(int, start_time_str.split(':'))
            start_time = india_tz.localize(datetime(today.year, today.month, today.day, hours, minutes))

            # ✅ If entered time already passed today → schedule for tomorrow
            if start_time < datetime.now(india_tz):
                start_time = start_time + timedelta(days=1)

        except Exception as e:
            print("⚠️ Invalid start time:", e)
            start_time = datetime.now(india_tz)
    else:
        start_time = datetime.now(india_tz)
    


    # ✅ Build quiz document
    quiz_doc = {
        "quiz_code": quiz_code,
        "Quiz_name": Quiz_name,
        "num_questions": num_questions,
        "timer_minutes": int(request.form.get('timer_minutes', 5)),
        "is_shuffled": request.form.get('is_shuffled') == 'true',
        "creator_token": creator_token,
        "usn_start": request.form.get('usn_start', '').upper(),
        "usn_end": request.form.get('usn_end', '').upper(),
        "start_time": start_time  # 🕒 timezone-aware and fixed
    }

    db.quizzes.insert_one(quiz_doc)

    print("✅ Quiz created successfully")
    print("🕒 Chosen Start Time (IST):", start_time.strftime("%Y-%m-%d %H:%M:%S %Z"))

    # ✅ Store questions
    quiz = db.quizzes.find_one({"quiz_code": quiz_code})
    quiz_id = quiz["_id"]

    for i in range(1, num_questions + 1):
        q_text = request.form.get(f'question_text_{i}', '').strip()
        options = [
            request.form.get(f'option_{i}_{j}', '').strip()
            for j in range(1, 5)
            if request.form.get(f'option_{i}_{j}', '').strip()
        ]
        correct_answer = request.form.get(f'correct_answer_{i}', '').strip()
        description = request.form.get(f'description_{i}', '').strip()

        db.questions.insert_one({
            "quiz_id": quiz_id,
            "question_text": q_text,
            "options": options,
            "correct_answer": correct_answer,
            "description": description
        })

    return render_template(
        'quiz_created.html',
        quiz_code=quiz_code,
        quiz_link=quiz_link,
        start_time=start_time.strftime("%I:%M %p")
    )

# --------------------------------------------------------------------------
# --- 5. Join, Attempt, and Result Views ---
# --------------------------------------------------------------------------
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/join', methods=['GET', 'POST'])
@app.route('/join', methods=['GET', 'POST'])
def join_quiz():
    db = get_db_connection()

    if request.method == 'POST':
        quiz_code = request.form['quiz_code'].upper()
        usn = request.form['usn'].upper()
        student_name = request.form['student_name']

        # Find the quiz by code
        quiz = db.quizzes.find_one({"quiz_code": quiz_code})
        if not quiz:
            return render_template('join_quiz.html', error="Invalid Quiz Code.", preset_code=quiz_code)

        # ✅ Check USN range restriction
        usn_start = quiz.get("usn_start")
        usn_end = quiz.get("usn_end")

        if usn_start and usn_end:
            # Ensure both are uppercase for comparison
            if not (usn_start <= usn <= usn_end):
                return render_template(
                    'join_quiz.html',
                    error=f"❌ You are not authorized to attempt this quiz. Allowed range: {usn_start} - {usn_end}",
                    preset_code=quiz_code
                )

        # Check if this user already attempted
        existing_attempt = db.attempts.find_one({"quiz_id": quiz["_id"], "usn": usn})
        if existing_attempt:
            return render_template('join_quiz.html', error="You already attempted this quiz.", preset_code=quiz_code)

        # Record a new attempt
        db.attempts.insert_one({
            "quiz_id": quiz["_id"],
            "usn": usn,
            "student_name": student_name,
            "score": 0,
            "start_time": datetime.now(),
            "end_time": None
        })

        # Save attempt info in session
        session['attempt_data'] = {
            'quiz_id': str(quiz["_id"]),
            'usn': usn,
            'student_name': student_name
        }

        return redirect(url_for('attempt_quiz'))

    preset_code = request.args.get('code', '')
    return render_template('join_quiz.html', preset_code=preset_code, error=None)





# --------------------------------------------------------------------------
# --- ATTEMPT QUIZ ROUTE (Updated with proper start_time waiting) ---
# --------------------------------------------------------------------------
@app.route('/quiz/attempt', methods=['GET', 'POST'])
def attempt_quiz():
    # --- [1] Verify session data ---
    if 'attempt_data' not in session:
        return redirect(url_for('join_quiz'))

    db = get_db_connection()
    attempt_data = session['attempt_data']

    # --- [2] Validate quiz_id ---
    try:
        quiz_id = ObjectId(attempt_data['quiz_id'])
    except Exception:
        return "Invalid quiz ID format", 400

    # --- [3] Fetch quiz and questions ---
    quiz = db.quizzes.find_one({"_id": quiz_id})
    print("📄 Raw start_time from MongoDB:", quiz.get("start_time"), type(quiz.get("start_time")))

    if not quiz:
        return "Quiz not found", 404

    questions = list(db.questions.find({"quiz_id": quiz_id}))
    if not questions:
        return "No questions found for this quiz.", 404

    # --- [4] Time-based access control (LIVE WAITING LOGIC) ---
    # --- 4. Check if current time < start_time ---
    india_tz = pytz.timezone("Asia/Kolkata")
    now = datetime.now(india_tz)

    quiz_start = quiz.get("start_time")

    # 🕒 Convert start_time safely from MongoDB
    if isinstance(quiz_start, str):
        try:
            quiz_start = datetime.fromisoformat(quiz_start)
        except Exception:
            try:
                quiz_start = datetime.strptime(quiz_start, "%Y-%m-%d %H:%M:%S")
            except Exception:
                quiz_start = None

    # Ensure timezone-awareness
    # ✅ Ensure quiz_start is timezone-aware and in IST
    if quiz_start:
        # MongoDB usually stores in UTC, so convert it to IST
        if quiz_start.tzinfo is None:
            quiz_start = pytz.UTC.localize(quiz_start)
        quiz_start = quiz_start.astimezone(india_tz)

        print("🕒 Converted quiz_start to IST:", quiz_start.strftime("%Y-%m-%d %H:%M:%S"))
        print("🕒 Current time (IST):", now.strftime("%Y-%m-%d %H:%M:%S"))


    # 🕐 Compare current time with quiz start
    if quiz_start and now < quiz_start:
        wait_seconds = int((quiz_start - now).total_seconds())
        print(f"⏳ Waiting: Quiz starts in {wait_seconds} seconds")
        print(f"🕒 Current Time (IST): {now.strftime('%H:%M:%S')}")
        print(f"🕒 Quiz Start Time (IST): {quiz_start.strftime('%H:%M:%S')}")
        display_start = quiz_start.astimezone(india_tz)
        return render_template("attempt_quiz.html", quiz=quiz, wait_seconds=wait_seconds, display_start=display_start)



    # --- [5] Clean and normalize question options ---
    import ast
    q_list = []
    for q in questions:
        q_dict = dict(q)
        opts = q_dict.get("options", [])

        if isinstance(opts, str):
            try:
                opts = ast.literal_eval(opts)
            except Exception:
                opts = [x.strip() for x in opts.split('|') if x.strip()]

        if not isinstance(opts, list):
            opts = [str(opts)]
        else:
            opts = [str(o).strip() for o in opts if o]

        while len(opts) < 4:
            opts.append(f"Option {chr(65 + len(opts))}")

        q_dict["options"] = opts
        q_list.append(q_dict)

    # --- [6] Shuffle if enabled ---
    if quiz.get("is_shuffled", False):
        random.shuffle(q_list)

    # --- [7] Handle submission ---
    if request.method == "POST":
        total_score = 0
        submitted_answers = {}

        for q in q_list:
            qid = str(q["_id"])
            submitted = request.form.get(f"q_{qid}")
            submitted_answers[qid] = submitted
            if submitted and submitted.strip() == q.get("correct_answer", "").strip():
                total_score += 1

        db.attempts.update_one(
            {"quiz_id": quiz_id, "usn": attempt_data["usn"]},
            {"$set": {
                "quiz_code": quiz.get("quiz_code"),
                "student_name": attempt_data.get("student_name"),
                "usn": attempt_data.get("usn"),
                "score": total_score,
                "submitted_answers": submitted_answers,
                "end_time": datetime.now(india_tz)
            }},
            upsert=True
        )

        session.pop("attempt_data", None)
        return render_template(
            "attempt_quiz.html",
            quiz=quiz,
            score=total_score,
            total_questions=len(q_list),
            questions=q_list,
            submitted=submitted_answers,
            final_view=True
        )

    # --- [8] Render quiz normally if GET ---
    return render_template("attempt_quiz.html", quiz=quiz, questions=q_list, final_view=False)




@app.route('/autosave', methods=['POST'])
def autosave():
    data = request.json
    quiz_id = ObjectId(data.get('quiz_id'))
    usn = data.get('usn')
    submitted = data.get('answers', {})

    db = get_db_connection()
    db.attempts.update_one(
        {"quiz_id": quiz_id, "usn": usn},
        {"$set": {"submitted_answers": submitted, "last_autosave": datetime.now()}},
        upsert=True
    )
    return jsonify({"status": "saved"})



@app.route('/results/<quiz_code>')
def view_results(quiz_code):
    token = request.args.get('token')
    if not token:
        return "Missing token", 403

    db = get_db_connection()
    quiz = db.quizzes.find_one({"quiz_code": quiz_code})

    if not quiz or quiz.get("creator_token") != token:
        return "Invalid quiz or token", 404

    # Fetch all attempts for this quiz, sorted by score descending
    attempts = list(db.attempts.find({"quiz_id": quiz["_id"]}).sort("score", -1))

    return render_template('result.html', quiz=quiz, attempts=attempts)

@app.route('/create_quiz_config')
def create_quiz_config():
    return render_template('create_quiz_config.html')

# --------------------------------------------------------------------------
# --- 6. NEW FEATURE: View Quiz (Student or Creator) ---
# --------------------------------------------------------------------------
@app.route('/view', methods=['GET', 'POST'])
def view_quiz():
    if request.method == 'POST':
        user_type = request.form.get('user_type')
        if user_type == 'student':
            return redirect(url_for('student_view'))
        elif user_type == 'creator':
            return redirect(url_for('creator_view'))
    return render_template('view_choice.html')

@app.route('/view/student', methods=['GET', 'POST'])
def student_view():
    if request.method == 'POST':
        usn = request.form.get('usn', '').upper().strip()
        db = get_db_connection()

        print("🔍 Searching attempts for USN:", usn)

        # ✅ Fetch all attempts for this USN
        attempts = list(db.attempts.find({"usn": usn}).sort("end_time", -1))
        print("🧾 Found attempts:", len(attempts))

        if not attempts:
            return render_template('student_view.html', error=f"No quiz attempts found for USN: {usn}")

        results = []
        for a in attempts:
            quiz = db.quizzes.find_one({"_id": a["quiz_id"]})
            if quiz:
                results.append({
                    "attempt_id": str(a["_id"]),  # for detail link
                    "quiz_code": quiz.get("quiz_code", "N/A"),
                    "quiz_name": quiz.get("Quiz_name", "Unknown"),
                    "score": a.get("score", 0),
                    "date": a.get("end_time").strftime("%Y-%m-%d") if a.get("end_time") else "N/A",
                    "time": a.get("end_time").strftime("%H:%M:%S") if a.get("end_time") else "N/A"
                })

        print("✅ Final Results:", results)
        return render_template('student_results.html', attempts=results, usn=usn)

    return render_template('student_view.html')

@app.route('/view/attempt/<attempt_id>')
def view_attempt_details(attempt_id):
    db = get_db_connection()

    # --- Find attempt document ---
    try:
        attempt = db.attempts.find_one({"_id": ObjectId(attempt_id)})
    except Exception as e:
        return f"Invalid attempt ID: {e}", 400

    if not attempt:
        return "Attempt not found", 404

    # --- Fetch quiz info ---
    quiz = db.quizzes.find_one({"_id": attempt["quiz_id"]})
    if not quiz:
        return "Quiz not found", 404

    # --- Fetch all quiz questions ---
    questions = list(db.questions.find({"quiz_id": quiz["_id"]}))
    if not questions:
        return "No questions found", 404

    # --- Extract answers ---
    submitted_answers = attempt.get("submitted_answers", {})

    detailed_list = []
    for q in questions:
        q_id = str(q["_id"])
        student_answer = submitted_answers.get(q_id, "Not Attempted")
        correct_answer = q.get("correct_answer", "N/A")
        description = q.get("description", "")
        q_text = q.get("question_text", "")
        options = q.get("options", [])

        detailed_list.append({
            "question": q_text,
            "options": options,
            "student_answer": student_answer,
            "correct_answer": correct_answer,
            "description": description,
            "is_correct": (student_answer == correct_answer)
        })

    return render_template(
        "attempt_details.html",
        quiz=quiz,
        attempt=attempt,
        detailed=detailed_list
    )


@app.route('/view/creator', methods=['GET', 'POST'])
def creator_view():
    if request.method == 'POST':
        # ✅ STEP 1: Get quiz_code and creator_token directly from form
        quiz_code = request.form.get('quiz_code', '').strip().upper()
        creator_token = request.form.get('creator_token', '').strip()

        if not quiz_code or not creator_token:
            return render_template('creator_view.html', error="Please enter both Quiz Code and Creator Token.")

        # ✅ STEP 2: Connect to MongoDB
        db = get_db_connection()

        # ✅ STEP 3: Validate quiz + token
        quiz = db.quizzes.find_one({"quiz_code": quiz_code, "creator_token": creator_token})
        if not quiz:
            return render_template('creator_view.html', error="Invalid Quiz Code or Creator Token.")

        # ✅ STEP 4: Fetch attempts
        attempts = list(db.attempts.find({"quiz_id": quiz["_id"]}))
        if not attempts:
            return render_template('creator_view.html', error="No attempts found for this quiz.")

        # ✅ STEP 5: Extract & format time fields
        result_data = []
        for a in attempts:
            usn = a.get("usn", "N/A")
            name = a.get("student_name", "Unknown")
            score = a.get("score", 0)
            start_time = a.get("start_time")
            end_time = a.get("end_time")

            def fmt_time(t):
                if isinstance(t, datetime):
                    return t.strftime("%H:%M:%S")
                elif isinstance(t, str) and len(t) >= 19:
                    return t[11:19]
                return "N/A"

            def fmt_date(t):
                if isinstance(t, datetime):
                    return t.strftime("%Y-%m-%d")
                elif isinstance(t, str) and len(t) >= 10:
                    return t[:10]
                return "N/A"

            date_str = fmt_date(start_time or end_time)
            start_str = fmt_time(start_time)
            end_str = fmt_time(end_time)

            result_data.append({
                "name": name,
                "usn": usn,
                "date": date_str,
                "start_time": start_str,
                "end_time": end_str,
                "score": score
            })

        total_attempts = len(result_data)

        # ✅ STEP 6: Export results if requested
        creator_email = request.form.get('creator_email')
        export_format = request.form.get('export_format')

        if export_format == 'pdf':
            return export_results_pdf(result_data, quiz_code, total_attempts)
        elif export_format == 'docx':
            return export_results_docx(result_data, quiz_code, total_attempts)

        if creator_email:
            try:
                send_results_docx_via_email(creator_email, result_data, quiz_code, total_attempts)
                print(f"📧 Results sent to {creator_email}")
            except Exception as e:
                print(f"⚠️ Failed to send email: {e}")

        # ✅ STEP 7: Render HTML table
        return render_template(
            'creator_view.html',
            quiz_code=quiz_code,
            results=result_data,
            total_attempts=total_attempts
        )

    # Default GET
    return render_template('creator_view.html')





def export_results_pdf(results, quiz_code, total_attempts):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, txt=f"Quiz Results - Code: {quiz_code}", ln=True, align="C")

    pdf.set_font("Arial", "", 12)
    pdf.cell(200, 10, txt=f"Total Attempts: {total_attempts}", ln=True, align="C")

    pdf.ln(10)
    pdf.set_font("Arial", "B", 11)
    pdf.cell(10, 10, "#", 1)
    pdf.cell(35, 10, "Date", 1)
    pdf.cell(30, 10, "Start", 1)
    pdf.cell(30, 10, "End", 1)
    pdf.cell(30, 10, "USN", 1)
    pdf.cell(35, 10, "Name", 1)
    pdf.cell(20, 10, "Score", 1, ln=True)

    pdf.set_font("Arial", "", 10)
    for i, r in enumerate(results, 1):
        pdf.cell(10, 10, str(i), 1)
        pdf.cell(35, 10, str(r.get("date", "") or ""), 1)
        pdf.cell(30, 10, str(r.get("start_time", "") or ""), 1)
        pdf.cell(30, 10, str(r.get("end_time", "") or ""), 1)
        pdf.cell(30, 10, str(r.get("usn", "") or ""), 1)
        pdf.cell(35, 10, str(r.get("name", "") or ""), 1)
        pdf.cell(20, 10, str(r.get("score", "") or ""), 1, ln=True)

    # ✅ Save to INTERNAL STORAGE (Documents folder)
    user_docs = os.path.join(os.path.expanduser("~"), "Documents", "Quiz_Results")
    os.makedirs(user_docs, exist_ok=True)

    filename = f"Quiz_{quiz_code}_Results_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    full_path = os.path.join(user_docs, filename)
    pdf.output(full_path)

    print(f"✅ PDF saved successfully in: {full_path}")

    # ✅ Optional: automatically open file after saving (Windows only)
    try:
        os.startfile(full_path)
    except Exception as e:
        print(f"⚠️ Could not auto-open file: {e}")

    # ✅ Show confirmation message in browser
    return render_template("creator_view.html", success=f"PDF stored successfully at: {full_path}")


def export_results_docx(results, quiz_code, total_attempts):
    doc = Document()
    doc.add_heading(f"Quiz Results - Code: {quiz_code}", 0)
    doc.add_paragraph(f"Total Attempts: {total_attempts}\n")

    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Date"
    hdr_cells[2].text = "Start Time"
    hdr_cells[3].text = "End Time"
    hdr_cells[4].text = "USN"
    hdr_cells[5].text = "Name"
    hdr_cells[6].text = "Score"

    for i, r in enumerate(results, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(r.get("date", "") or "")
        row_cells[2].text = str(r.get("start_time", "") or "")
        row_cells[3].text = str(r.get("end_time", "") or "")
        row_cells[4].text = str(r.get("usn", "") or "")
        row_cells[5].text = str(r.get("name", "") or "")
        row_cells[6].text = str(r.get("score", "") or "")

    # ✅ SAVE TO INTERNAL STORAGE (Documents folder)
    user_docs = os.path.join(os.path.expanduser("~"), "Documents", "Quiz_Results")
    os.makedirs(user_docs, exist_ok=True)

    filename = f"Quiz_{quiz_code}_Results_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    full_path = os.path.join(user_docs, filename)

    doc.save(full_path)

    print(f"✅ DOCX saved successfully in: {full_path}")

    # Optional: automatically open file (Windows only)
    try:
        os.startfile(full_path)
    except:
        pass

    return render_template("creator_view.html", success=f"DOCX stored successfully at: {full_path}")


@app.route('/db_test')
def db_test():
    try:
        db = get_db_connection()
        db.command("ping")
        return "✅ MongoDB connected successfully"
    except Exception as e:
        return f"❌ MongoDB connection failed: {e}"




@app.route('/debug/questions/<quiz_code>')
def debug_questions(quiz_code):
    db = get_db_connection()
    quiz = db.quizzes.find_one({"quiz_code": quiz_code.upper()})
    if not quiz:
        return "Quiz not found", 404

    questions = list(db.questions.find({"quiz_id": quiz["_id"]}))
    for q in questions:
        print("🧩", q["question_text"], q["options"])
    return jsonify([
        {"question": q["question_text"], "options": q["options"]}
        for q in questions
    ])


def send_results_docx_via_email(email, results, quiz_code, total_attempts):
    from io import BytesIO

    # ✅ Step 1: Generate DOCX in memory
    doc = Document()
    doc.add_heading(f"Quiz Results - Code: {quiz_code}", 0)
    doc.add_paragraph(f"Total Attempts: {total_attempts}\n")

    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Date"
    hdr_cells[2].text = "Start Time"
    hdr_cells[3].text = "End Time"
    hdr_cells[4].text = "USN"
    hdr_cells[5].text = "Name"
    hdr_cells[6].text = "Score"

    for i, r in enumerate(results, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(r.get("date", "") or "")
        row[2].text = str(r.get("start_time", "") or "")
        row[3].text = str(r.get("end_time", "") or "")
        row[4].text = str(r.get("usn", "") or "")
        row[5].text = str(r.get("name", "") or "")
        row[6].text = str(r.get("score", "") or "")

    # Save DOCX to memory (not disk)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # ✅ Step 2: Create and send email
    msg = Message(
        subject=f"📘 Quiz Results Report - {quiz_code}",
        recipients=[email],
        body=f"""
Hello,

Please find attached the full quiz results for Quiz Code: {quiz_code}.
Total Attempts: {total_attempts}

- Quiz System
"""
    )

    # Attach DOCX
    msg.attach(
        filename=f"Quiz_{quiz_code}_Results.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buffer.read()
    )

    mail.send(msg)
    print(f"📧 DOCX report sent successfully to {email}")


# --------------------------------------------------------------------------
# --- 7. Run the app ---
# --------------------------------------------------------------------------
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)
